import sys
import subprocess

# Ensure python-docx is installed
try:
    import docx
except ImportError:
    print("python-docx not found. Installing...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    """Sets background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets cell margins (padding)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_document():
    doc = Document()
    
    # Configure page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Styles Setup
    style_normal = doc.styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'Segoe UI'
    font_normal.size = Pt(11)
    font_normal.color.rgb = RGBColor(0x33, 0x41, 0x55) # Sleek slate gray
    
    # Custom Title Style
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("PROMPTPILOT TACTICAL INSTRUCTION")
    title_run.font.name = 'Segoe UI'
    title_run.font.size = Pt(26)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB) # PromptPilot Theme Blue
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Hugging Face API Authentication & Key Setup Guide")
    subtitle_run.font.name = 'Segoe UI'
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.italic = True
    subtitle_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B) # Slate gray
    
    doc.add_paragraph() # Spacer
    
    # Warning/Status Box
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell = table.cell(0, 0)
    set_cell_background(cell, "FEF2F2") # Soft red background for warning
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    p = cell.paragraphs[0]
    p_run = p.add_run("⚠️ SYSTEM AUTHENTICATION EXPIRED / INSUFFICIENT PERMISSIONS")
    p_run.font.bold = True
    p_run.font.size = Pt(11)
    p_run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26) # Bold red
    
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(6)
    p_desc = p2.add_run(
        "A connection check was performed on the existing HF_TOKEN in your .env.local file. "
        "The server returned an HTTP 403 Forbidden error: 'This authentication method does not have "
        "sufficient permissions to call Inference Providers on behalf of user'."
    )
    p_desc.font.size = Pt(10.5)
    
    p3 = cell.add_paragraph()
    p3.paragraph_format.space_before = Pt(6)
    p_desc2 = p3.add_run(
        "Follow the steps below to generate a new token with the correct permissions and re-establish connection."
    )
    p_desc2.font.size = Pt(10.5)
    
    doc.add_paragraph() # Spacer
    
    # Section 1
    h1 = doc.add_paragraph()
    h1_run = h1.add_run("Step-by-Step Setup Guide")
    h1_run.font.size = Pt(18)
    h1_run.font.bold = True
    h1_run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A) # Dark slate
    
    steps = [
        ("Step 1: Access Hugging Face", "Open your web browser and go to the official Hugging Face website: https://huggingface.co. Log in to your account, or sign up for a new account if you do not have one."),
        ("Step 2: Navigate to Settings", "Click on your profile picture in the top-right corner of the Hugging Face homepage and select Settings from the dropdown menu."),
        ("Step 3: Access Access Tokens Tab", "In the left-hand navigation sidebar, click on Access Tokens (or go directly to: https://huggingface.co/settings/tokens)."),
        ("Step 4: Create a New Token", "Click on the Create new token button located on the top right of the tokens section."),
        ("Step 5: Select Token Type and Permissions (CRITICAL)", 
         "1. Under Token type, choose Fine-grained (recommended for security).\n"
         "2. Enter a Token name (e.g., PromptPilot-Inference).\n"
         "3. Under Permissions, scroll down to the Inference section and check the boxes for:\n"
         "   - Make calls to the serverless Inference API\n"
         "   - Make calls to Inference Providers (Required for the Hugging Face Router API)\n"
         "4. (Optional) If you choose a Classic token, make sure to select the Write scope."),
        ("Step 6: Copy Your Generated Token", "Click the Generate token button at the bottom of the page. Once created, copy the token. It should start with the prefix hf_."),
        ("Step 7: Update Environment Configuration", "Open the .env.local file in the root of your PromptPilot workspace (c:\\Users\\Welcome\\Downloads\\Promptpilot-main\\Promptpilot\\.env.local). Find the line starting with HF_TOKEN= and replace the existing key with your new token:\n\nHF_TOKEN=hf_yourNewTokenHere\n\nSave the file."),
        ("Step 8: Restart Next.js Server", "If your local Next.js development server is currently running, stop it (Ctrl+C in terminal) and start it again with npm run dev so Next.js reads the updated environment variable.")
    ]
    
    for title, desc in steps:
        p_title = doc.add_paragraph()
        p_title.paragraph_format.space_before = Pt(12)
        p_title.paragraph_format.space_after = Pt(4)
        run_title = p_title.add_run(f"■ {title}")
        run_title.font.bold = True
        run_title.font.size = Pt(12)
        run_title.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
        
        # Split description by newlines to render separate paragraphs beautifully
        lines = desc.split('\n')
        for line in lines:
            if not line.strip():
                continue
            p_desc = doc.add_paragraph()
            p_desc.paragraph_format.left_indent = Inches(0.25)
            p_desc.paragraph_format.space_after = Pt(4)
            run_desc = p_desc.add_run(line)
            run_desc.font.size = Pt(10.5)
        
    # Verification
    h2 = doc.add_paragraph()
    h2_run = h2.add_run("How to Verify the Connection")
    h2_run.font.size = Pt(16)
    h2_run.font.bold = True
    h2_run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    
    vp = doc.add_paragraph()
    vp_run = vp.add_run(
        "To verify that your token is fully functional with the Hugging Face router, run the following verification "
        "command in the terminal of the PromptPilot workspace root directory:\n\n"
        "npx tsx test-router.ts\n\n"
        "If successful, the script will output the AI's response model name and message payload indicating that "
        "connectivity is successfully established."
    )
    vp_run.font.size = Pt(10.5)
    
    doc.save("HuggingFace_API_Setup_Instructions.docx")
    print("Document created successfully at HuggingFace_API_Setup_Instructions.docx")

if __name__ == "__main__":
    create_document()
